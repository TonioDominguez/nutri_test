import streamlit as st
import pandas as pd
import ast  # To turn string lists into real lists Para convertir las listas almacenadas como strings en listas reales
from docx import Document # To doc export part 4 the menu
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to load the CSV data
@st.cache_data
def load_data():
    data = pd.read_csv("recetas_cat_0301.csv")  # Reemplaza con el nombre de tu archivo
    # Convertir las listas almacenadas como strings en listas reales
    data["apartat"] = data["apartat"].apply(ast.literal_eval)
    return data

@st.cache_data #Function to load the ingredients for the shopping list
def load_ingredients_data():
    return pd.read_csv("ingredientes_estandarizados.csv")

# Function to generate the weekly menu
def generate_menu(selected_dishes, data):
    menu = {}
    days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
    meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

    for day in days:
        menu[day] = {}
        for meal in meals:
            dishes = selected_dishes.get(f"{day}-{meal}", [])
            if dishes:
                # Obtain details of each plate
                dish_details = [
                    {
                        "nom": dish,
                        "ingredients": data[data["nom"] == dish].iloc[0]["ingredients"],
                        "passos": data[data["nom"] == dish].iloc[0]["passos"]
                    }
                    for dish in dishes
                ]
                menu[day][meal] = dish_details
            else:
                menu[day][meal] = []
    return menu

# Function to generate the shopping list

def generate_shopping_list(menu, ingredients_data):
    shopping_list = {}

    # Recorrer el menú
    for day, meals in menu.items():
        for meal, dishes in meals.items():
            if dishes:
                for dish in dishes:
                    dish_name = dish['nom']
                    # Buscar los ingredientes estandarizados para esta receta
                    dish_ingredients = ingredients_data[ingredients_data['receta'] == dish_name]

                    for _, ingredient in dish_ingredients.iterrows():
                        ing_name = ingredient['ingrediente_estandar']
                        cantidad = ingredient['cantidad']
                        unidad = ingredient['unidad']

                        if ing_name in shopping_list:
                            if pd.notna(cantidad):
                                shopping_list[ing_name]['cantidad'] += cantidad
                        else:
                            shopping_list[ing_name] = {
                                'cantidad': cantidad if pd.notna(cantidad) else 0,
                                'unidad': unidad if pd.notna(unidad) else ''
                            }

    return shopping_list

#Function to generate weekly meny GRID

def create_menu_grid(menu):
    """Crea una tabla/grid con el menú semanal"""
    days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
    meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

    # Ahora filas = días, columnas = tipos de comida
    menu_grid = pd.DataFrame(index=days, columns=meals)

    for day in days:
        for meal in meals:
            if day in menu and meal in menu[day]:
                dishes = menu[day][meal]
                if dishes:
                    menu_grid.loc[day, meal] = ", ".join([dish["nom"] for dish in dishes])
                else:
                    menu_grid.loc[day, meal] = "———"
            else:
                menu_grid.loc[day, meal] = "———"
    return menu_grid

# Function to set borderlines to grid in WORD
def set_table_borders(table):
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Tipo de borde (línea simple)
                border.set(qn('w:sz'), '4')       # Grosor del borde
                border.set(qn('w:space'), '0')    # Espaciado
                border.set(qn('w:color'), '000000')  # Color del borde (negro)
                tcBorders.append(border)

            tcPr.append(tcBorders)

# Load the data
st.title("Generador de Menús Setmanals")
st.sidebar.title("Opcions de filtratge")
data = load_data()

# Mapping meal times to CSV columns
meal_columns = {
    "Desdejuni": "desdejuni",
    "Dinar": "dinar",
    "Sopar": "sopar",
    "Snack": "snack"
}

# Sidebar filters
st.sidebar.subheader("Filtrar per moment del dia")
moment_of_day = st.sidebar.selectbox(
    "Selecciona moment del dia",
    options=["Tots"] + list(meal_columns.keys())
)

# Extract unique categories
unique_categories = sorted({category for sublist in data["apartat"] for category in sublist})

st.sidebar.subheader("Filtrar per categoria única")
selected_categories = st.sidebar.multiselect(
    "Selecciona una o més categories",
    options= unique_categories,
    default=[],
    placeholder="Selecciona una o més opcions"
)

# Apply filters
filtered_data = data.copy()

# Filter by moment of the day
if moment_of_day != "Tots":
    filtered_data = filtered_data[filtered_data[meal_columns[moment_of_day]] == 1]

# Filter by multiple categories
if selected_categories:
    filtered_data = filtered_data[filtered_data["apartat"].apply(
        lambda x: all(category in x for category in selected_categories)
    )]

# Display filtered dishes
st.subheader("Plats disponibles")
st.dataframe(filtered_data)

# Initialize session state for selected dishes
if "selected_dishes" not in st.session_state:
    st.session_state.selected_dishes = {}

# Weekly menu selection
st.subheader("Selecciona els plats per a cada dia i moment")
days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

for day in days:
    st.write(f"### {day}")
    for meal in meals:
        # Create a unique key for each selection
        key = f"{day}-{meal}"
        # Get the previously selected value or default
        previous_selection = st.session_state.selected_dishes.get(key, [])
        # Show only dishes corresponding to the selected moment of the day
        options = list(filtered_data["nom"].unique())
        options_with_previous = list(set(options + previous_selection))
        # Update the session state with the new selection
        st.session_state.selected_dishes[key] = st.multiselect(
            f"{meal} ({day})", 
            options_with_previous, 
            default=previous_selection,
            placeholder="Selecciona una o més opcions"
        )

# Generate the weekly menu
if st.button("Generar Menú Setmanal"):
    menu = generate_menu(st.session_state.selected_dishes, data)  # Pass 'data' as the second argument
    ingredients_data = load_ingredients_data()
    shopping_list = generate_shopping_list(menu, ingredients_data)

    #Grid visualization
    
    st.subheader("Vista de Taula del Menú Setmanal")
    menu_grid = create_menu_grid(menu)
    st.table(menu_grid)

    #Shopping list visualization
    st.subheader("Llista de la compra")
    for ingredient in sorted(shopping_list.keys()):
        details = shopping_list[ingredient]
        cantidad = details['cantidad']
        unidad = details['unidad']
        if cantidad > 0:
            st.write(f"- {cantidad} {unidad} {ingredient}")
        else:
            st.write(f"- {ingredient}")
    
    #Menu visualization
    st.subheader("Menú Setmanal")
    for day, meals in menu.items():
        st.write(f"### **{day}**")
        for meal, dishes in meals.items():
            st.write(f"**{meal}:**")
            if dishes:
                for dish in dishes:
                    st.write(f"- {dish['nom']}")
                    with st.expander(f"Detalls de {dish['nom']}"):
                        st.write("**Ingredients:**")
                        st.write(", ".join(ast.literal_eval(dish["ingredients"])))
                        st.write("**Passos:**")
                        for step in ast.literal_eval(dish["passos"]):
                            st.write(f"- {step}")
            else:
                st.write("- No seleccionat")


                        

                        
    # Create Menu DOC
    

    doc = Document()
    doc.add_heading("Menú Setmanal", level=1)

    
    # Grid in DOC
    
    table = doc.add_table(rows=len(days)+1, cols=len(meals)+1)

    table.cell(0, 0).text = ""
    for i, meal in enumerate(meals):
        table.cell(0, i+1).text = meal
    for i, day in enumerate(days):
        table.cell(i+1, 0).text = day

    for i, day in enumerate(days):
        for j, meal in enumerate(meals):
            dishes = menu[day][meal]
            if not dishes:
                table.cell(i + 1, j + 1).text = "———"
            else:
                table.cell(i + 1, j + 1).text = ", ".join([dish["nom"] for dish in dishes])

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1  
    
    set_table_borders(table) #Function to apply borders

    # Create shopping list in DOC
    doc.add_page_break()
    doc.add_heading("Llista de la compra", level=1)

    # Ordenar la lista de la compra alfabéticamente
    for ingredient in sorted(shopping_list.keys()):
        details = shopping_list[ingredient]
        cantidad = details['cantidad']
        unidad = details['unidad']
        if cantidad > 0:
            doc.add_paragraph(f"- {cantidad} {unidad} {ingredient}")
        else:
            doc.add_paragraph(f"- {ingredient}")
    
    # Corpus DOC
    
    for day, meals in menu.items():
        doc.add_heading(day, level=2)
        for meal, dishes in meals.items():
            doc.add_heading(meal, level=3)
            if dishes:
                for dish in dishes:
                    p = doc.add_paragraph()
                    run = p.add_run("Plat:")
                    run.bold = True
                    p.add_run(f" {dish['nom']}")

                    # Ingredients list bullets style
                    p = doc.add_paragraph()
                    run = p.add_run("Ingredients:")
                    run.bold = True
                    ingredients_list = ast.literal_eval(dish["ingredients"])
                    for ingredient in ingredients_list:
                        p = doc.add_paragraph(ingredient, style='ListBullet')

                    # Steps list bullets style
                    p = doc.add_paragraph()
                    run = p.add_run("Passos:")
                    run.bold = True
                    steps_list = ast.literal_eval(dish["passos"])
                    for i, step in enumerate(steps_list, 1):
                        p = doc.add_paragraph(f"{i}. {step}")

            else:
                doc.add_paragraph("No seleccionat")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Export DOC
    st.download_button(
        label="Descarregar Menú en Word",
        data=buffer,
        file_name="menu_setmanal.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )