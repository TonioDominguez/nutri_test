import streamlit as st
import pandas as pd
import ast  # To turn string lists into real lists Para convertir las listas almacenadas como strings en listas reales
from docx import Document # To doc export part 4 the menu

# Function to load the CSV data
@st.cache_data
def load_data():
    data = pd.read_csv("recetas_cat_0301.csv")  # Reemplaza con el nombre de tu archivo
    # Convertir las listas almacenadas como strings en listas reales
    data["apartat"] = data["apartat"].apply(ast.literal_eval)
    return data

# Function to generate the weekly menu
def generate_menu(selected_dishes, data):
    menu = {}
    days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
    meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

    for day in days:
        menu[day] = {}
        for meal in meals:
            dish_name = selected_dishes.get(f"{day}-{meal}", "No seleccionat")
            if dish_name != "No seleccionat":
                # Get the full details of the selected dish
                dish_details = data[data["nom"] == dish_name].iloc[0]
                menu[day][meal] = {
                    "nom": dish_name,
                    "ingredients": dish_details["ingredients"],
                    "passos": dish_details["passos"]
                }
            else:
                menu[day][meal] = {"nom": "No seleccionat", "ingredients": "", "passos": ""}
    return menu

#Function to generate weekly meny GRID

def create_menu_grid(menu):
    """Crea una tabla/grid con el menú semanal"""
    # Crear un DataFrame para la visualización en grid
    days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
    meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

    menu_grid = pd.DataFrame(index=meals, columns=days)

    #for day in days:
     #   for meal in meals:
      #      if day in menu and meal in menu[day]:
       #         menu_grid.loc[meal, day] = menu[day][meal]["nom"]
        #    else:
         #       menu_grid.loc[meal, day] = "Lliure"

    for day in days:
        for meal in meals:
            if day in menu and meal in menu[day]:
                if menu[day][meal]["nom"] == "No seleccionat":
                    menu_grid.loc[meal, day] = "———"  
                else:
                    menu_grid.loc[meal, day] = menu[day][meal]["nom"]
            else:
                menu_grid.loc[meal, day] = "———"         
    
    
    return menu_grid

# Load the data
st.title("Generador de Menús Setmanals")
st.sidebar.title("Opcions de filtratge")
data = load_data()

# Mapping meal times to CSV columns
meal_columns = {
    "Desdejuni": "desdejuni",
    "Dinar": "dinar",
    "Sopar": "sopar",
    "Snack": "snack"
}

# Sidebar filters
st.sidebar.subheader("Filtrar per moment del dia")
moment_of_day = st.sidebar.selectbox(
    "Selecciona moment del dia",
    options=["Tots"] + list(meal_columns.keys())
)

# Extract unique categories
unique_categories = sorted({category for sublist in data["apartat"] for category in sublist})

st.sidebar.subheader("Filtrar per categoria única")
selected_categories = st.sidebar.multiselect(
    "Selecciona una o més categories",
    options= unique_categories,
    default=[]
)

# Apply filters
filtered_data = data.copy()

# Filter by moment of the day
if moment_of_day != "Tots":
    filtered_data = filtered_data[filtered_data[meal_columns[moment_of_day]] == 1]

# Filter by multiple categories
if selected_categories:
    filtered_data = filtered_data[filtered_data["apartat"].apply(
        lambda x: all(category in x for category in selected_categories)
    )]

# Display filtered dishes
st.subheader("Plats disponibles")
st.dataframe(filtered_data)

# Initialize session state for selected dishes
if "selected_dishes" not in st.session_state:
    st.session_state.selected_dishes = {}

# Weekly menu selection
st.subheader("Selecciona els plats per a cada dia i moment")
days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

for day in days:
    st.write(f"### {day}")
    for meal in meals:
        # Create a unique key for each selection
        key = f"{day}-{meal}"
        # Get the previously selected value or default to "No seleccionat"
        previous_selection = st.session_state.selected_dishes.get(key, "No seleccionat")
        # Show only dishes corresponding to the selected moment of the day
        options = ["No seleccionat"] + list(filtered_data["nom"].unique())
        # If the previous selection is not in the current options, keep it
        if previous_selection not in options and previous_selection != "No seleccionat":
            options.append(previous_selection)
        # Update the session state with the new selection
        st.session_state.selected_dishes[key] = st.selectbox(
            f"{meal} ({day})", options, index=options.index(previous_selection) if previous_selection in options else 0
        )

# Generate the weekly menu
if st.button("Generar Menú Setmanal"):
    menu = generate_menu(st.session_state.selected_dishes, data)  # Pass 'data' as the second argument
    
    #Grid visualization
    
    st.subheader("Vista de Taula del Menú Setmanal")
    menu_grid = create_menu_grid(menu)
    st.table(menu_grid)
    
    #Menu visualization
    st.subheader("Menú Setmanal")
    for day, meals in menu.items():
        st.write(f"### **{day}**")
        for meal, details in meals.items():
            st.write(f"**{meal}:** {details['nom']}")
            if details["nom"] != "No seleccionat":
                with st.expander(f"Detalls de {details['nom']}"):
                    st.write("**Ingredients:**")
                    st.write(", ".join(ast.literal_eval(details["ingredients"])))
                    st.write("**Passos:**")
                    for step in ast.literal_eval(details["passos"]):
                        st.write(f"- {step}")


                        

                        
    # Create Menu DOC
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("Menú Setmanal", level=1)

    
    # Grid in DOC
    
    table = doc.add_table(rows=len(meals)+1, cols=len(days)+1)

    table.cell(0, 0).text = ""
    for i, day in enumerate(days):
        table.cell(0, i+1).text = day
    for i, meal in enumerate(meals):
        table.cell(i+1, 0).text = meal

    for i, meal in enumerate(meals):
        for j, day in enumerate(days):
            if menu[day][meal]["nom"] == "No seleccionat":
                table.cell(i+1, j+1).text = "———" 
            else:
                table.cell(i+1, j+1).text = menu[day][meal]["nom"]

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1  
    
    
    # Corpus DOC
    
    for day, meals in menu.items():
        doc.add_heading(day, level=2)
        for meal, details in meals.items():
            doc.add_heading(meal, level=3)
            doc.add_paragraph(f"**Plat:** {details['nom']}")
            if details["nom"] != "No seleccionat":
                doc.add_paragraph("**Ingredients:**")
                doc.add_paragraph(", ".join(ast.literal_eval(details["ingredients"])))
                doc.add_paragraph("**Passos:**")
                for step in ast.literal_eval(details["passos"]):
                    doc.add_paragraph(f"- {step}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Export DOC
    st.download_button(
        label="Descarregar Menú en Word",
        data=buffer,
        file_name="menu_setmanal.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )