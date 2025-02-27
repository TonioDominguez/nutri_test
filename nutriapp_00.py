import streamlit as st
import pandas as pd
import ast  # To turn string lists into real lists Para convertir las listas almacenadas como strings en listas reales
from docx import Document # To doc export part 4 the menu
from io import BytesIO

# MySQL connection config
def get_database_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="tomate160121",
        database="nutriapp"
    )

# Function to load the data
@st.cache_data
def load_data():
    try:
        conn = get_database_connection()
        query = """
        SELECT nom, ingredients, passos, desdejuni, dinar, sopar, snack, apartat
        FROM recetas
        """
        data = pd.read_sql(query, conn)
        conn.close()

        # Change column text that contains lists
        
        data["apartat"] = data["apartat"].apply(ast.literal_eval)
        data["ingredients"] = data["ingredients"].apply(ast.literal_eval)
        data["passos"] = data["passos"].apply(ast.literal_eval)

        return data
    except Exception as e:
        st.error(f"Error al conectar con la base de datos: {e}")
        return None

# Function to generate the weekly menu
def generate_menu(selected_dishes, data):
    menu = {}
    days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
    meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

    for day in days:
        menu[day] = {}
        for meal in meals:
            dish_name = selected_dishes.get(f"{day}-{meal}", "No seleccionat")
            if dish_name != "No seleccionat":
                # Get the full details of the selected dish
                dish_details = data[data["nom"] == dish_name].iloc[0]
                menu[day][meal] = {
                    "nom": dish_name,
                    "ingredients": dish_details["ingredients"],
                    "passos": dish_details["passos"]
                }
            else:
                menu[day][meal] = {"nom": "No seleccionat", "ingredients": "", "passos": ""}
    return menu

# Load the data
st.title("Generador de Menús Setmanals")
st.sidebar.title("Opcions de filtratge")
data = load_data()

# Mapping meal times to CSV columns
meal_columns = {
    "Desdejuni": "desdejuni",
    "Dinar": "dinar",
    "Sopar": "sopar",
    "Snack": "snack"
}

# Sidebar filters
st.sidebar.subheader("Filtrar per moment del dia")
moment_of_day = st.sidebar.selectbox(
    "Selecciona moment del dia",
    options=["Tots"] + list(meal_columns.keys())
)

# Extract unique categories
unique_categories = sorted({category for sublist in data["apartat"] for category in sublist})

st.sidebar.subheader("Filtrar per categoria única")
selected_categories = st.sidebar.multiselect(
    "Selecciona una o més categories",
    options= unique_categories,
    default=[]
)

# Apply filters
filtered_data = data.copy()

# Filter by moment of the day
if moment_of_day != "Tots":
    filtered_data = filtered_data[filtered_data[meal_columns[moment_of_day]] == 1]

# Filter by multiple categories
if selected_categories:
    filtered_data = filtered_data[filtered_data["apartat"].apply(
        lambda x: all(category in x for category in selected_categories)
    )]

# Display filtered dishes
st.subheader("Plats disponibles")
st.dataframe(filtered_data)

# Initialize session state for selected dishes
if "selected_dishes" not in st.session_state:
    st.session_state.selected_dishes = {}

# Weekly menu selection
st.subheader("Selecciona els plats per a cada dia i moment")
days = ["Dilluns", "Dimarts", "Dimecres", "Dijous", "Divendres", "Dissabte", "Diumenge"]
meals = ["Desdejuni", "Dinar", "Sopar", "Snack"]

for day in days:
    st.write(f"### {day}")
    for meal in meals:
        # Create a unique key for each selection
        key = f"{day}-{meal}"
        # Get the previously selected value or default to "No seleccionat"
        previous_selection = st.session_state.selected_dishes.get(key, "No seleccionat")
        # Show only dishes corresponding to the selected moment of the day
        options = ["No seleccionat"] + list(filtered_data["nom"].unique())
        # If the previous selection is not in the current options, keep it
        if previous_selection not in options and previous_selection != "No seleccionat":
            options.append(previous_selection)
        # Update the session state with the new selection
        st.session_state.selected_dishes[key] = st.selectbox(
            f"{meal} ({day})", options, index=options.index(previous_selection) if previous_selection in options else 0
        )

# Generate the weekly menu
if st.button("Generar Menú Setmanal"):
    menu = generate_menu(st.session_state.selected_dishes, data)  # Pass 'data' as the second argument
    st.subheader("Menú Setmanal")
    for day, meals in menu.items():
        st.write(f"### **{day}**")
        for meal, details in meals.items():
            st.write(f"**{meal}:** {details['nom']}")
            if details["nom"] != "No seleccionat":
                with st.expander(f"Detalls de {details['nom']}"):
                    st.write("**Ingredients:**")
                    st.write(", ".join(ast.literal_eval(details["ingredients"])))
                    st.write("**Passos:**")
                    for step in ast.literal_eval(details["passos"]):
                        st.write(f"- {step}")

    # Export the menu to Word

    doc = Document()
    doc.add_heading("Menú Setmanal", level=1)

    for day, meals in menu.items():
        doc.add_heading(day, level=2)
        for meal, details in meals.items():
            doc.add_heading(meal, level=3)
            doc.add_paragraph(f"**Plat:** {details['nom']}")
            if details["nom"] != "No seleccionat":
                doc.add_paragraph("**Ingredients:**")
                doc.add_paragraph(", ".join(ast.literal_eval(details["ingredients"])))
                doc.add_paragraph("**Passos:**")
                for step in ast.literal_eval(details["passos"]):
                    doc.add_paragraph(f"- {step}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Descarregar Menú en Word",
        data=buffer,
        file_name="menu_setmanal.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )